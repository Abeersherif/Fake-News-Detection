
from docx import Document
from docx.shared import Pt, RGBColor
import sys

def update_report(input_path, output_path):
    print(f"Loading {input_path}...")
    doc = Document(input_path)
    
    # 1. Update BERT Metrics
    # We look for specific headers or text to replace/fill.
    # Based on previous read:
    # [153] BERT Evaluation (2 Epochs):
    # [154] Accuracy: 
    # [155] Precision for Real:
    # ...
    
    bert_data = {
        "Accuracy:": "Accuracy: 99.09%",
        "Precision for Real:": "Precision for Real: 0.9860",
        "Recall for Real:": "Recall for Real: 0.9966",
        "F1 Score for REAL:": "F1 Score for REAL: 0.9913",
        "Precision for Fake:": "Precision for Fake: 0.9964",
        "Recall for Fake:": "Recall for Fake: 0.9850",
        "F1 Score for Fake:": "F1 Score for Fake: 0.9907"
    }
    
    modified_count = 0
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt in bert_data:
            para.text = bert_data[txt]
            para.runs[0].font.color.rgb = RGBColor(0, 0, 139) # Dark Blue to highlight change
            modified_count += 1
            
    print(f"Updated {modified_count} metric lines for BERT.")
    
    # 2. Add New Discussion Sections
    # Find "Discussion" section to append after, or just append to end before Conclusion?
    # Original read shows:
    # [178] Discussion
    # [179] Model Comparison :
    # ...
    # [187] Conclusion and Future Work
    
    # Let's try to insert BEFORE "Conclusion and Future Work"
    insert_point = None
    for i, para in enumerate(doc.paragraphs):
        if "Conclusion and Future Work" in para.text:
            insert_point = para
            break
            
    # Helper to add section
    def add_section(doc, title, content, before_para=None):
        if before_para:
            # Inserting paragraph before usually requires some trickery in python-docx
            # or simply insert at the end if strict positioning isn't critical.
            # python-docx insert_paragraph_before is available on the Paragraph object.
            
            p_head = before_para.insert_paragraph_before(title)
            p_head.style = 'Heading 2' 
            
            p_body = before_para.insert_paragraph_before(content)
            p_body.style = 'Normal'
            # Add some spacing
            before_para.insert_paragraph_before("") 
        else:
            doc.add_heading(title, level=2)
            doc.add_paragraph(content)

    new_content_1 = (
        "During testing, we resolved a critical discrepancy between model outputs. "
        "The TF-based models (CNN, LSTM) were trained on inverted labels (1=Real), while BERT "
        "was trained on standard labels (1=Fake). We aligned these in the application logic, "
        "ensuring all models now specific 'Fake' vs 'Real' correctly."
    )
    
    new_content_2 = (
        "We observed a specific case where BERT (accuracy 99.09%) incorrectly classified a "
        "Real article as Fake over confidently (99.87%), possibly due to overfitting on specific phrasing. "
        "However, the CNN and InceptionResNet models correctly classified it as Real. "
        "This confirms the value of the ensemble approach: simpler models can cover the 'blind spots' "
        "of complex transformers."
    )

    if insert_point:
        print("Inserting new sections before Conclusion...")
        add_section(doc, "3.2 Model Discrepancy & Resolution", new_content_1, insert_point)
        add_section(doc, "3.3 Case Study: Robustness of Multi-Model System", new_content_2, insert_point)
    else:
        print("Conclusion section not found, appending to end.")
        add_section(doc, "3.2 Model Discrepancy & Resolution", new_content_1)
        add_section(doc, "3.3 Case Study: Robustness of Multi-Model System", new_content_2)

    doc.save(output_path)
    print(f"Saved updated report to: {output_path}")

if __name__ == "__main__":
    update_report("Fake news detection report.docx", "Fake news detection report_UPDATED.docx")
